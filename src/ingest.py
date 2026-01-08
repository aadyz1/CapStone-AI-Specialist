import os
import tempfile
import subprocess
from typing import List
from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_chroma import Chroma
from langchain_core.documents import Document as LCDocument
from docx import Document as DocxDocument

from src.config import CHROMA_DIR, JD_COLLECTION, RESUME_COLLECTION, OPENAI_API_KEY



def _load_txt(path: str) -> List[LCDocument]:
    """Load a .txt file into LangChain Document objects."""
    loader = TextLoader(path, encoding="utf-8")
    return loader.load()


def _load_docx(path: str) -> List[LCDocument]:
    """Load a .docx file using python-docx and return one LangChain Document."""
    doc = DocxDocument(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs).strip()
    return [LCDocument(page_content=text, metadata={"source_file": os.path.basename(path)})]


def _load_pdf(path: str) -> List[LCDocument]:
    """
    Load a .pdf file using PyPDFLoader (requires pypdf).

    PyPDFLoader returns a list of Documents (often one per page).
    We keep them as-is and chunk afterwards.
    """
    loader = PyPDFLoader(path)
    return loader.load()


def _convert_doc_to_docx(doc_path: str) -> str:
    """
    Convert legacy .doc files to .docx using LibreOffice (soffice).

    Notes:
    - .doc is the old Microsoft Word binary format.
    - python-docx cannot read .doc directly.
    - If LibreOffice is installed, we can convert .doc -> .docx automatically.
    """
    out_dir = tempfile.mkdtemp(prefix="doc_to_docx_")
    cmd = ["soffice", "--headless", "--convert-to", "docx", "--outdir", out_dir, doc_path]

    try:
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    except FileNotFoundError as e:
        raise RuntimeError(
            "This file is a .doc (old Word format). Python cannot read .doc directly. "
            "Install LibreOffice so the 'soffice' command is available OR convert .doc to .docx manually."
        ) from e
    except subprocess.CalledProcessError as e:
        raise RuntimeError(
            "LibreOffice was found but failed to convert the .doc file. Please convert to .docx manually."
        ) from e

    base = os.path.splitext(os.path.basename(doc_path))[0]
    converted = os.path.join(out_dir, base + ".docx")

    if not os.path.exists(converted):
        docx_files = [os.path.join(out_dir, f) for f in os.listdir(out_dir) if f.lower().endswith(".docx")]
        if not docx_files:
            raise RuntimeError("Conversion finished but no .docx file was produced. Please convert manually.")
        converted = docx_files[0]

    return converted


def load_document(path: str) -> List[LCDocument]:
    """
    Unified loader supporting:
    - .txt
    - .docx
    - .doc  (auto-convert to .docx if LibreOffice exists)
    - .pdf
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == ".txt":
        return _load_txt(path)
    if ext == ".docx":
        return _load_docx(path)
    if ext == ".doc":
        converted = _convert_doc_to_docx(path)
        return _load_docx(converted)
    if ext == ".pdf":
        return _load_pdf(path)

    raise ValueError(f"Unsupported file type: {ext}. Use .txt, .docx, .doc, or .pdf")


def chunk_documents(docs: List[LCDocument]) -> List[LCDocument]:
    """Split long text into smaller chunks for better vector retrieval."""
    splitter = RecursiveCharacterTextSplitter(
    chunk_size=800, 
    chunk_overlap=150)
    return splitter.split_documents(docs)


def build_vectorstores():
    """Create/load persistent Chroma vector stores from disk."""
    embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)

    jd_vs = Chroma(
        collection_name=JD_COLLECTION,
        embedding_function=embeddings,
        persist_directory=CHROMA_DIR,
    )

    resume_vs = Chroma(
        collection_name=RESUME_COLLECTION,
        embedding_function=embeddings,
        persist_directory=CHROMA_DIR,
    )

    return jd_vs, resume_vs


def ingest_all(jd_path: str, resumes_dir: str):
    """
    End-to-end ingestion:
    - Load JD (.txt/.docx/.doc/.pdf)
    - Load resumes (.txt/.docx/.doc/.pdf)
    - Chunk text
    - Embed + store in Chroma (with metadata)
    """
    jd_vs, resume_vs = build_vectorstores()

    # ---- Ingest JD ----
    jd_docs = load_document(jd_path)
    jd_chunks = chunk_documents(jd_docs)
    for c in jd_chunks:
        c.metadata["doc_type"] = "jd"

    jd_vs.add_documents(jd_chunks)
    ###########jd_vs.persist()

    # ---- Ingest resumes ----
    for filename in os.listdir(resumes_dir):
        lower = filename.lower()
        if not (lower.endswith(".txt") or lower.endswith(".docx") or lower.endswith(".doc") or lower.endswith(".pdf")):
            continue

        candidate_id = os.path.splitext(filename)[0]
        path = os.path.join(resumes_dir, filename)

        resume_docs = load_document(path)
        resume_chunks = chunk_documents(resume_docs)

        for c in resume_chunks:
            c.metadata["doc_type"] = "resume"
            c.metadata["candidate_id"] = candidate_id
            c.metadata["source_file"] = filename

        resume_vs.add_documents(resume_chunks)

    #######resume_vs.persist()
    return True

